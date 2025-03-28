# from pandas.core.frame import StringIO
from io import BytesIO
import streamlit as st
import os
from supabase import create_client, Client
from dotenv import load_dotenv
import supabase
from pypdf import PdfReader
# import textract # textract 1.6.5 requires six~=1.12.0, but you have six 1.16.0 which is incompatible.
from docx import Document
from st_supabase_connection import SupabaseConnection
import argon2
from openai import OpenAI
import time
import re
import json
from streamlit_cookie_banner import cookie_banner

st.set_page_config(
    page_title="TestMaker",
    page_icon="🧑‍🏫",
)

load_dotenv()

api_key: str = os.getenv('SUPABASE_KEY')
supabase_url: str = os.getenv('SUPABASE_URL')
supabase: Client = create_client(supabase_url, api_key)
llm_key = os.getenv("OPENROUTER_API_KEY")
llm_base_url: str = os.getenv("OPENROUTER_BASE_URL")
model = "google/gemini-2.0-flash-exp:free"
llm = OpenAI(
    base_url=llm_base_url,
    api_key=llm_key
)

if "lecture_text" not in st.session_state:
    st.session_state.lecture_text = ""
    
# prompt0 = """You are a school teacher who needs to generate 5 tests in the affirmative form with 4 possible answers and one correct answer.
# The tests must meet the following requirements:
# • formulate the text of the assignment in an affirmative form, at the end of the sentence there is a colon sign - ":"!
# The wording of the text of each assignment must be in an affirmative form, WITHOUT the "question" sign (?) and WITHOUT question words!
# the very wording of the task should not contain hints on the correct answer!
# • place keywords in the text of the assignment at the beginning of the sentence;
# • You cannot use the answer options "All of the above", "All of the above options", "All of the above", "All except", "All answer options are correct", "Both options", "All answer options are incorrect";
# • all answer options should be equally attractive: similar in both appearance and grammatical structure, the correct answer should not contain a grammatical hint;
# After the answer options, one correct answer option is indicated for each question.
# - exclude lengthy arguments, repetitions, and complex syntactic ones in the text
# turns, double negation, as well as the words "sometimes", "often", "always", "all",
# "never";
# - exclude the words: "specify", "select", "list", "name", "all of
# the listed", "all except";
# - the answers to the task should be meaningful, similar both in
# appearance and grammatical structure, and attractive to choose from;

# - every wrong answer must be plausible, credible and convincing;
# - there should be no hints on the correct answer in the text of the task

# Before the correct answer, it should be written - "The correct answer"
# After each correct question there should be two empty lines "\n\n"
# Before generating new questions , write - 'Here are the corrected questions:'
# The questions should correspond to this context {}. 5 Generated questions with correct answers, after each question the correct answer is indicated:"""
    
first_prompt = """Ты школьный учитель, которому необходимо сгенерировать 5 тестов в утвердительной форме с 4 вариантами ответов и одним верным вариантом ответа.
Тесты должны соответствовать следующим требованиям:
•	формулировать текст задания в утвердительной форме , в конце предложения стоит знак двоеточия - ":"!
Формулировка текста каждого задания должна быть в утвердительной форме, БЕЗ знака "вопрос" (?) и БЕЗ вопросительных слов!
в самой формулировке задания нет подсказок на верный вариант ответа!
•	ключевые слова в тексте задания размещать в начале предложения;
•	Нельзя использовать варианты ответов "Все вышеперечисленные","Все вышеперечисленные варианты", «Все перечисленное», «Все, кроме», «Все варианты ответов верны», "Оба варианта", «Все варианты ответов неверны»;
•	все варианты ответов должны быть одинаково привлекательными: похожими как по внешнему виду, так и по грамматической структуре, правильный вариант ответа не должен содержать грамматической подсказки;
после вариантов ответов для каждого вопроса указан один верный вариант ответа.

- исключать в тексте пространные рассуждения, повторы, сложные синтаксические
обороты, двойное отрицание, а также – слова «иногда», «часто», «всегда», «все»,
«никогда»;
- исключать слова: «укажите», «выберите», «перечислите», «назовите», «все из
перечисленных», «все, кроме»;
- варианты ответа на задание должны быть содержательными, похожими как по
внешнему виду, так и по грамматической структуре, привлекательными для выбора;

- каждый неправильный вариант ответа должен быть правдоподобным, внушающим доверие и убедительным;
- в тексте задания не должно быть подсказок на верный вариант ответа

Перед верным вариантом ответа должно быть написано - "Верный ответ"
После каждого верного вопроса должно быть две пустые строки "\n\n"
Перед генерацией новых вопросов , напиши - 'Вот исправленные вопросы:'
Вопросы должны соответствовать этому контексту {}. 

Пример теста 1:
4. Основной принцип цифровой этики включает в себя:
А) Применение активного слушания
Б) Помнить про безопасность
В) Уважение частной жизни
Г) Честность с партнером

Верный ответ: В) Уважение частной жизни

Пример теста 2:
Хорошее информационное сообщение должно обладать комбинацией свойств, включая:
А) Полезность и ценность
Б) Достоверность и объективность
В) Объективность и ценность
Г) Точность и достоверность

Верный ответ: Б) Достоверность и объективность

5 Сгенериованных вопросов с верным вариантом ответа, после каждого вопроса указан верный ответ:"""


second_prompt = """Даны сгенерированные ранее тобой вопросы {}, те вопросы в которых последний вариант ответа говорит о верности всех вышеупомянутых по типу 'все вышеперечисленное' необходимо изменить, поскольку такой вариант ответа не приемлем. Должен быть только один верный вариант ответа. Таким образом, необходимо верные вопросы оставить без изменения, а те, которые не соответствуют требованиям заменить верными формулировками вариантов ответов. 
Тесты должны соответствовать следующим требованиям:
•  формулировать текст задания в утвердительной форме , в конце предложения стоит знак двоеточия - ":"!
Формулировка текста каждого задания должна быть в утвердительной форме, БЕЗ знака "вопрос" (?) и БЕЗ вопросительных слов!
в самой формулировке задания не должно быть подсказок на верный вариант ответа!
•  ключевые слова в тексте задания размещать в начале предложения;
•  Нельзя использовать варианты ответов "Все вышеперечисленные","Все вышеперечисленные варианты", «Все перечисленное», «Все, кроме», «Все варианты ответов верны», "Оба варианта", «Все варианты ответов неверны»;
•  все варианты ответов должны быть одинаково привлекательными: похожими как по внешнему виду, так и по грамматической структуре, правильный вариант ответа не должен содержать грамматической подсказки;
после вариантов ответов для каждого вопроса указан один верный вариант ответа.
Предоставь сгенерированные вопросы в формате JSON. JSON должен иметь следующую структуру: [{{'question': '1. ...', 'choices': ['A ...', 'B ...' ,], 'correct_answer': '...'}}, ...]. Здесь 'question' - формулировка самого задания, 'choices' - это варианты ответов, 'correct_answer' - это верный вариант ответа, который также содержится в 'choices'.
Убедись, что JSON правильно отформатирован. Также не забудь добавить разделитель ',' после каждого варианта в 'choices'.
Не добавляйте никаких комментариев до или после создания тестов, просто создайте тесты.
 """

# prompt2 = """You are a professor with expertise in every possible field and should create an exam on the topic of the Input PDF. "
# "Using the attached lecture slides (please analyze thoroughly), create a Master-level multiple-choice exam. The exam should contain multiple-choice and single-choice questions, "
# "appropriately marked so that students know how many options to select. Create 5 realistic exam questions covering the entire content. Provide the output in a JSON format. "
# "The JSON should have the structure: [{'question': '1. ...', 'choices': ['A ...', 'B ...',], 'correct_answer': '...', 'explanation': '...'}, ...]. Ensure the JSON is valid and properly formatted. Also, don't forget to add the ',' delimiter after each last option in the 'choices' so that it can be parsed correctly as JSON."""

def login_form(
    *,
    title: str = "Log In/Sign Up",
    user_tablename : str = "users",
    username_col: str = "username",
    email_col: str = "email",
    password_col: str = "password",
    signup_title: str = "Create new account",
    login_title: str = "Login to existing account",
    allow_signup: bool = True,
) -> Client:
    """Create a user login/signup form with email validation and password hashing which is then saved to supabase db."""
    
    client = st.connection(name = "supabase", type = SupabaseConnection)
    password_hasher = argon2.PasswordHasher()
    
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
        
    if "username" not in st.session_state:
        st.session_state.username = None
        
    if not st.session_state["authenticated"]:
        with st.expander("Authentication", expanded = True):
            if allow_signup:
                signup_tab, login_tab = st.tabs(["New User? Sign Up", "Existing User? Log In"])
            else:
                login_tab = st.container()
            if allow_signup:
                with signup_tab:
                    with st.form(key = "signup"):
                        username = st.text_input(label = "Enter a username")
                        user_email = st.text_input(label = "Enter your email")
                        password = st.text_input(label = "Create password", help = "Password cannot be recovered if lost", type = "password")
                        confirmed_password = st.text_input(label = "Confirm Password", type = "password")
                        assert password == confirmed_password # todo: add more context
                        hashed_password = password_hasher.hash(password)
                        if st.form_submit_button(label="Sign Up", type = "primary"):
                            if "@" not in user_email: # todo: switch to more robust email validation library.
                                st.error("Invalid Email address")
                                st.stop()
                            try:
                                client.table(user_tablename).insert({username_col: username, email_col: user_email, password_col: hashed_password}).execute()
                            except Exception as e:
                                st.error(e.message)
                            else:
                                st.session_state["authenticated"] = True
                                st.session_state["username"] = username
                                st.success("Sign up successful")
                                st.rerun()
                                
            with login_tab:
                with st.form(key = "login"):
                    username = st.text_input(label = "Username:")
                    password = st.text_input(label = "Password", type = "password")
                    
                    if st.form_submit_button(label = "Log In", type = "primary"):
                        response = client.table(user_tablename).select(f"{username_col}, {password_col}").eq(username_col, username).execute()

                        if len(response.data) > 0:
                            db_password = response.data[0]["password"]
                            if password_hasher.verify(db_password, password):
                                st.session_state["authenticated"] = True
                                st.session_state["username"] = username
                                st.success("Log in successful")
                                st.rerun()
                            else:
                                st.error("Incorrect Password")
                        else:
                            st.error("Username or Password incorrect")
    return client


def signout(client):
    # client.auth.sign_out()
    st.session_state["authenticated"] = False
    st.session_state["username"] = None
    st.session_state.lecture_text = None

def clean_filename(filename):
    base, ext = os.path.splitext(filename)
    base_eng = re.match('^[a-zA-Z0-9._-]+$', base)
    if not base_eng:
        toutf8 = base.encode('utf-8')
        hexenc = toutf8.hex()
        return f'{hexenc}{ext}'
    cleaned_base = re.sub(r'[^\w\s-]', '', base).strip().replace(' ', '_')
    return f'{cleaned_base}{ext}'
    
def upload_file(file, client):
    
    if file is not None:
        file_type = file.type
        file_name = clean_filename(file.name)
        file_contents = file.getvalue()
        
        if file_type == "text/plain":
            st.session_state.lecture_text = file.read().decode('utf-8')
            client.upload(bucket_id = "t3files", source = "local", file = file, destination_path = f"lectures/{file_name}", overwrite = "true")
            file_db_path = f"{supabase_url}/storage/v1/object/public/t3files/lectures/{file_name}"
            response = client.table("users").select("username, id").eq("username", st.session_state.username).execute()
            user_id = response.data[0]["id"]
            client.table("lectures").insert(dict(file_path = file_db_path, user_id = user_id)).execute()
    
        elif file_type == "application/pdf":
            reader = PdfReader(file)
            lecture_text = ""
            for page in reader.pages:
                lecture_text += page.extract_text()
            st.session_state.lecture_text = lecture_text
            supabase.storage.from_("t3files").upload(
                path = f"lectures/{file_name}",
                file = file_contents,
                
                file_options = {
                    "content-type": file_type,
                    "x-upsert": "true"
                }
            )
            file_db_path = f"{supabase_url}/storage/v1/object/public/t3files/lectures/{file_name}"
            response = client.table("users").select("username, id").eq("username", st.session_state.username).execute()
            user_id = response.data[0]["id"]
            client.table("lectures").insert(dict(file_path = file_db_path, user_id = user_id)).execute()
        
        # elif file_type == "application/msword":
            # st.session_state.text = textract.process(file) # might not work due to dependency incompatibility. waiting for textract to merge dependency update.
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            st.session_state.lecture_text = "\n".join([para.text for para in doc.paragraphs])
            supabase.storage.from_("t3files").upload(
                path = f"lectures/{file_name}",
                file = file_contents,
                
                file_options = {
                    "content-type": file_type,
                    "x-upsert": "true"
                }
            )
            file_db_path = f"{supabase_url}/storage/v1/object/public/t3files/lectures/{file_name}"
            response = client.table("users").select("username, id").eq("username", st.session_state.username).execute()
            user_id = response.data[0]["id"]
            client.table("lectures").insert(dict(file_path = file_db_path, user_id = user_id)).execute()
        
        else:
            st.error("Unsupported file type!")
        
    return st.session_state.lecture_text

@st.fragment
def chunk_text(text, max_tokens = 15000):
    sentences = text.split('. ')
    chunks = []
    chunk = sentences[0] + '. '
    
    for sentence in sentences[1:]:
        if len(chunk) + len(sentence) + 2 > max_tokens:
            chunks.append(chunk)
            chunk = sentence + ". "
        else:
            chunk += sentence + ". "
    if chunk:
        chunks.append(chunk)
    return chunks


def generate_test(prompt, text):
    test_questions = []
    # chunked_text = chunk_text(text)
    for i in range(len(text)):
    
        combined_message = prompt.format(text[i])
        # print(combined_message ,'\n')
        # print('-' * 50)
        messages = [
            {
                "role": "user",
                "content": combined_message
            }
        ]
        responses = llm.chat.completions.create(
            model = model,
            messages = messages,
            stream = True,
            temperature = 0.5
        )
        test = ""
        for response in responses:
            test += response.choices[0].delta.content or ""
            
        # print(test)
        test_questions.append(test)
    # print(test_questions)
    
    return test_questions
  
def parse_generated_test(test):
    full_test = []
    # print(f'test: {test} \n\n')
    for i in range(len(test)):
        try:
            json_start = test[i].find('[')
            json_end = test[i].rfind(']') + 1
            json_str = test[i][json_start:json_end]
            # print(f'\n\njson string: {json_str}\n\n')
            questions = json.loads(json_str)
            full_test.extend(questions)
        except json.JSONDecodeError as e:
            st.error(f"JSON parsing error: {e}")
            st.error("Response from OpenAI:")
            st.text(test)
            return None
    # print(f'full_test: {full_test}\n\n')
    return full_test
    # print(questions)
        
def display_questions(questions):
    # print(questions)
    for index, question in enumerate(questions):
        test_question = question["question"]
        st.write(f"{index+1}: {test_question}")
        st.write("выбор:")
        for i, choice in enumerate(question["choices"]):
            st.write(f"{choice}")
        st.write(f"верный ответ: {question['correct_answer']}")
        
        feedback = submit_feedback(index)

feedbacks = {}

@st.fragment
def submit_feedback(q_index):
    feedback = st.feedback(key = q_index)
    if feedback is not None:
        feedbacks[f"Q{q_index+1}"] = feedback
    # print(feedbacks)
    return feedbacks
    
@st.fragment
def add_comment():
    with st.form("feedback comment", clear_on_submit=True):
        comment = st.text_input("Как можно улучшить тест? Расскажите как можно подробнее")
        st.form_submit_button("Отправить", use_container_width=True)
        return comment

@st.fragment
def save_test_to_db(questions, comment):
    user = supabase.table("users").select("username, id").eq("username", st.session_state.username).execute()
    user_id = user.data[0]["id"]
    lecture = supabase.table("lectures").select("id, user_id").eq("user_id", user_id).execute()
    lecture_id = lecture.data[0]["id"]
    supabase.table("tests").insert(dict(user_id = user_id, lecture_id = lecture_id, test_text = questions, test_feedback = feedbacks, test_comment = comment)).execute()

@st.fragment
def download_test(questions):
    doc = Document()
    text = ''
    for index, question in enumerate(questions):
        text += f'{index+1}: {question['question'][2:]}'
        text += '\n'
        for i, choice in enumerate(question['choices']):
            text += choice
            text += '\n'
        text += f"\nверный ответ: {question['correct_answer']}"
        text += '\n\n'
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    q_doc = buf.read()
    
    comment = add_comment()
    
    if st.download_button("Скачать тест", data = q_doc, mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type = "primary", use_container_width=True):
        save_test_to_db(questions, comment) 
        

def main():
    client = login_form()
    
    if st.session_state["authenticated"]:
        consent = cookie_banner(
            banner_text="Мы используем файлы cookie, чтобы обеспечить вам наилучший сервис.\n",
            display=True,  # Set to False if you want to hide the banner
            link_text="Learn more",  # Optional: Link text (e.g., 'Learn more')
            link_url="https://disk.yandex.ru/i/94HD1oBJCWRGdA",  # Optional: URL to your privacy policy
            key="cookie_banner"  # Optional: A unique key to control re-rendering
        )
        
        if consent:
            st.context.cookies
            
        with st.sidebar:
            st.title("TestMaker")
            st.write(f"Привет 👋 {st.session_state.username}!")
            st.button(label = "Выход", on_click=signout, args=[client], type = "primary", use_container_width=True)
        
        with st.popover("Инструкция по использованию", use_container_width=True):
            st.write("""
                Уважаемый пользователь!
            Наш сервис для создания тестовых заданий находится на этапе бетта-тестирования. Мы будем благодарны за использование сервиса и за вашу обратную связь.\n
            Инструкция по использованию 
            1. Выберите текст лекции и загрузите его в поле Drag and drop file here. 
            2. Дождись загрузки лекции на платформу (знак вверху страницы – Running = загрузка). Когда знак загрузки исчез, переходите к следующему этапу. 
            3. Для генерации тестовых заданий нажмите кнопку – Сгенерировать тест. 
            4. Далее проставьте отметки нравится или не нравится по каждому тесту. Это поможет улучшить наш сервис. 
            5. Нажмите кнопку - Отправьте отзыв и скачайте ваши тесты. 
            6. В поле для ввода введите комментарий с обратной связью. 
            7. Нажмите кнопку "Отправить" для загрузки тестов в систему. 
            8. Для скачивания тестов нажми кнопку – Скачать тест""")
        
        file = st.file_uploader("Выберите или перетащите файл для загрузки", type=['txt', 'pdf', 'docx'])
        if file is not None:
            upload_file(file, client)
        # st.button("Загрузить файл", on_click=upload_file, args=[file, client], use_container_width=True)

        if st.button("Сгенерировать тест", use_container_width=True):
            # upload_file(file, client)
            chunked_text = chunk_text(st.session_state.lecture_text)
            # for i, chunk in enumerate(chunked_text):
            #     print(f"Chunk {i+1}: \n {chunk}\n")
            first_gen_text = generate_test(first_prompt, chunked_text)
            time.sleep(15)
            second_gen_text = generate_test(second_prompt, first_gen_text)
            # st.write(second_gen_text)
            parsed_test = parse_generated_test(second_gen_text)
            display_questions(parsed_test)
            with st.popover("Отправьте отзыв и скачайте", use_container_width=True):
                st.write("Оцените вопросы с помощью ':material/thumb_up:', если вам нравится, и ':material/thumb_down:', если вам не нравится")
                download_test(parsed_test)
        
        
        
if __name__ == "__main__":
    main()



